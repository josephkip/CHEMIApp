from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style Setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.name = 'Calibri'

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            rc[i].text = str(val)
            for p in rc[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def tip(text):
    p = doc.add_paragraph()
    r = p.add_run('💡 Tip: ')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run(text)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run('⚠ Note: ')
    r.bold = True
    r.font.color.rgb = RGBColor(200, 120, 0)
    p.add_run(text)

def steps(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f'{i}. {item}')
        p.paragraph_format.left_indent = Cm(1)

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MORERAN CHEMIST')
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Pharmacy Management System')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('USER GUIDE')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Version 1.0 — May 2026').font.size = Pt(12)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Getting Started',
    '   2.1 Logging In',
    '   2.2 Dashboard Overview',
    '   2.3 Changing Your Password',
    '   2.4 Updating Your Profile',
    '3. Point of Sale (Making Sales)',
    '   3.1 Creating a New Sale',
    '   3.2 Viewing Sales History',
    '   3.3 Printing a Receipt',
    '4. Inventory Management',
    '   4.1 Browsing Inventory',
    '   4.2 Adding a New Item',
    '   4.3 Editing an Item',
    '   4.4 Restocking an Item',
    '   4.5 Deleting an Item',
    '   4.6 Stock Alerts',
    '5. Categories',
    '6. Procurement',
    '   6.1 Managing Suppliers',
    '   6.2 Creating a Purchase Order (LPO)',
    '   6.3 Receiving Goods (GRN)',
    '7. Stock Take',
    '   7.1 Starting a Stock Take',
    '   7.2 Entering Physical Counts',
    '   7.3 Completing a Stock Take',
    '8. Analytics & Reports',
    '9. User Management (Super Admin)',
    '   9.1 Creating a New User',
    '   9.2 Editing Users',
    '   9.3 Disabling / Enabling Users',
    '   9.4 Resetting Passwords',
    '10. Role Permissions Reference',
    '11. Troubleshooting',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'Moreran Chemist is a complete pharmacy management system designed to streamline '
    'daily operations including sales processing, inventory tracking, procurement, '
    'stock auditing, and business analytics. The system is accessible via desktop '
    'application (Windows), mobile app (Android), and web browser.'
)
doc.add_paragraph(
    'This guide walks you through every function of the system, step by step, '
    'so you can confidently operate it from day one.'
)

doc.add_heading('System Roles', level=2)
doc.add_paragraph('The system has three user roles with different access levels:')
add_table(
    ['Role', 'Description', 'Key Abilities'],
    [
        ['Sales Attendant', 'Front-counter cashier', 'Process sales, view inventory, view alerts'],
        ['Admin', 'Pharmacy manager', 'All of the above + manage items, procurement, stock takes, analytics'],
        ['Super Admin', 'System owner', 'All of the above + create/manage user accounts'],
    ]
)

# ============================================================
# 2. GETTING STARTED
# ============================================================
doc.add_heading('2. Getting Started', level=1)

doc.add_heading('2.1 Logging In', level=2)
doc.add_paragraph('To access the system, you need a username and password provided by the Super Admin.')
doc.add_heading('Steps:', level=3)
steps([
    'Open the Moreran Chemist application (desktop app, mobile app, or web browser).',
    'Enter your Username (or email address) in the login field.',
    'Enter your Password.',
    'Click the "Login" button.',
    'You will be taken to the Dashboard if your credentials are correct.',
])
note('If your account has been disabled by the Super Admin, you will not be able to log in. Contact your administrator.')
tip('Your session will expire after a period of inactivity. You will need to log in again.')

doc.add_heading('2.2 Dashboard Overview', level=2)
doc.add_paragraph(
    'After logging in, you will see the main Dashboard. What you see depends on your role:'
)
bullets([
    "Sales Attendants: Quick access to the POS (sales) screen, inventory search, and stock alerts.",
    "Admins & Super Admins: Full dashboard with today's revenue, profit, sales count, low stock warnings, expiring items, and quick-access navigation to all modules.",
])

doc.add_heading('2.3 Changing Your Password', level=2)
doc.add_paragraph('All users can change their own password at any time.')
steps([
    'Navigate to your Profile or Settings page.',
    'Enter your Current Password.',
    'Enter your New Password.',
    'Confirm the new password.',
    'Click "Change Password".',
])
tip('Use a strong password with at least 8 characters, mixing letters, numbers, and symbols.')

doc.add_heading('2.4 Updating Your Profile', level=2)
steps([
    'Navigate to your Profile page.',
    'Edit your Full Name or Username as needed.',
    'Click "Save" or "Update Profile".',
    'The system will issue a new session token with your updated information.',
])

doc.add_page_break()

# ============================================================
# 3. POINT OF SALE
# ============================================================
doc.add_heading('3. Point of Sale (Making Sales)', level=1)
doc.add_paragraph('Available to: All roles (Sales Attendant, Admin, Super Admin)')

doc.add_heading('3.1 Creating a New Sale', level=2)
steps([
    'Navigate to the Sales or POS section.',
    'Search for the product by name or scan the barcode.',
    'Select the item from the results.',
    'Enter the quantity the customer is purchasing.',
    'The system automatically calculates the line total (price × quantity).',
    'Repeat steps 2–5 for additional items.',
    'Review the cart — verify all items and quantities are correct.',
    'Select the Payment Method: Cash, M-Pesa, Card, or Insurance.',
    '(Optional) Enter the customer name if needed.',
    'Click "Complete Sale" to finalize.',
    'The receipt will be generated automatically.',
])

doc.add_heading('What happens behind the scenes:', level=3)
bullets([
    'Stock is automatically deducted for each item sold.',
    'A stock movement record is created (type: "sale").',
    'Profit is calculated per item: (selling price − buying price) × quantity.',
    'A unique receipt number is generated.',
])

note('If an item does not have enough stock, the system will block the sale and display: "Insufficient stock for [Item Name]. Available: [X]".')

doc.add_heading('3.2 Viewing Sales History', level=2)
steps([
    'Navigate to the Sales History section.',
    'Use the filters to narrow results: date range, payment method, or search by receipt number / customer name.',
    'Click on any sale to see its full details (items, quantities, prices, cashier name).',
])

doc.add_heading('3.3 Printing a Receipt', level=2)
steps([
    'Open the sale details from Sales History.',
    'Click "Print Receipt" or "View Receipt".',
    'The receipt includes: pharmacy name, receipt number, date, items, totals, and payment method.',
])

doc.add_page_break()

# ============================================================
# 4. INVENTORY MANAGEMENT
# ============================================================
doc.add_heading('4. Inventory Management', level=1)

doc.add_heading('4.1 Browsing Inventory', level=2)
doc.add_paragraph('Available to: All roles')
doc.add_paragraph('Navigate to the Inventory section to see all active products.')
doc.add_heading('You can:', level=3)
bullets([
    'Search by product name, batch number, or supplier.',
    'Filter by category.',
    'Filter by stock status: OK (healthy stock), Low (at or below reorder level), Out (zero stock).',
    'Sort by name, price, stock quantity, or any column.',
    'Click on any item to see full details including stock movement history.',
])

doc.add_heading('4.2 Adding a New Item', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
steps([
    'Navigate to Inventory and click "Add Item" or the "+" button.',
    'Fill in the required fields:',
    '   • Name — the product name (e.g., "Paracetamol 500mg")',
    '   • Category — select from existing categories',
    '   • Buying Price — what you paid for the item',
    '   • Profit Margin (%) — the system auto-calculates the selling price',
    'Fill in optional fields as needed:',
    '   • Initial Stock Quantity (defaults to 0)',
    '   • Reorder Level — the threshold for low-stock alerts (defaults to 10)',
    '   • Expiry Date',
    '   • Batch Number',
    '   • Supplier',
    '   • Unit (e.g., pcs, boxes, strips)',
    '   • Description',
    'Click "Save" or "Add Item".',
])
tip('If you enter an initial stock quantity, the system automatically creates a stock movement record labeled "Initial stock".')

doc.add_heading('4.3 Editing an Item', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
steps([
    'Find the item in the Inventory list.',
    'Click on the item to open its details.',
    'Click "Edit".',
    'Modify any field (name, price, reorder level, etc.).',
    'Click "Save".',
])

doc.add_heading('4.4 Restocking an Item', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
steps([
    'Find the item in the Inventory list.',
    'Click "Restock" (or open the item and click the restock button).',
    'Enter the quantity to add.',
    '(Optional) Add notes describing the restock reason.',
    'Click "Confirm".',
])
doc.add_paragraph('The system will:')
bullets([
    'Add the quantity to the current stock.',
    'Create a stock movement record (type: "restock") with before/after quantities.',
])

doc.add_heading('4.5 Deleting an Item', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
doc.add_paragraph(
    'Deleting an item performs a "soft delete" — the item is marked as inactive and hidden '
    'from the inventory list, but its historical data (sales, movements) is preserved.'
)
steps([
    'Find the item and open its details.',
    'Click "Delete".',
    'Confirm the deletion when prompted.',
])

doc.add_heading('4.6 Stock Alerts', level=2)
doc.add_paragraph('Available to: All roles')
add_table(
    ['Alert Type', 'Condition', 'Action Needed'],
    [
        ['Low Stock', 'Stock quantity ≤ reorder level', 'Restock or create a purchase order'],
        ['Expiring Soon (Warning)', 'Expiry date within 30 days', 'Prioritize selling or arrange return'],
        ['Expiring Soon (Danger)', 'Expiry date within 7 days', 'Remove from shelf if necessary'],
    ]
)

doc.add_page_break()

# ============================================================
# 5. CATEGORIES
# ============================================================
doc.add_heading('5. Categories', level=1)
doc.add_paragraph('Categories organize your inventory (e.g., "Antibiotics", "Pain Relief", "Vitamins").')

doc.add_heading('Viewing Categories', level=2)
doc.add_paragraph('Available to: All roles — Navigate to the Categories section to see the list.')

doc.add_heading('Creating a Category', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
steps([
    'Click "Add Category".',
    'Enter the Category Name (required).',
    '(Optional) Enter a Description.',
    'Click "Save".',
])

doc.add_heading('Editing a Category', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
steps([
    'Click on the category to edit.',
    'Update the name or description.',
    'Click "Save".',
])

doc.add_page_break()

# ============================================================
# 6. PROCUREMENT
# ============================================================
doc.add_heading('6. Procurement', level=1)
doc.add_paragraph(
    'The Procurement module manages the full cycle of purchasing stock from suppliers: '
    'managing suppliers, creating purchase orders (LPOs), and recording goods received (GRNs).'
)

doc.add_heading('6.1 Managing Suppliers', level=2)
doc.add_heading('Viewing Suppliers', level=3)
doc.add_paragraph('Available to: All roles — see the list of registered suppliers.')

doc.add_heading('Adding / Editing a Supplier', level=3)
doc.add_paragraph('Available to: Admin and Super Admin only')
steps([
    'Navigate to Procurement → Suppliers.',
    'Click "Add Supplier" (or click an existing supplier to edit).',
    'Enter the supplier details (name, contact, etc.).',
    'Click "Save".',
])

doc.add_heading('6.2 Creating a Purchase Order (LPO)', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
doc.add_paragraph('A Purchase Order (LPO) is a formal record of what you intend to order from a supplier.')
steps([
    'Navigate to Procurement → Purchase Orders.',
    'Click "Create LPO" or "New Purchase Order".',
    'Select the Supplier from the dropdown.',
    'Add items to the order: select each item, enter quantity and buying price.',
    'The system automatically calculates the total order amount.',
    '(Optional) Add notes.',
    'Click "Submit" or "Create".',
])
note('Creating an LPO does NOT change your stock levels. It is only a record of what you plan to purchase.')

doc.add_heading('6.3 Receiving Goods (GRN — Goods Received Note)', level=2)
doc.add_paragraph('Available to: Admin and Super Admin only')
doc.add_paragraph(
    'When a delivery arrives from a supplier, you create a GRN to officially receive the goods '
    'into your inventory.'
)
steps([
    'Navigate to Procurement → Goods Received.',
    'Click "Create GRN" or "Receive Goods".',
    '(Optional) Link to an existing LPO if this delivery fulfils a purchase order.',
    'Select the Supplier.',
    'For each item received, enter: the item, quantity received, and buying price.',
    '(Optional) Add receiving notes.',
    'Click "Submit" or "Create".',
])

doc.add_heading('What happens when you create a GRN:', level=3)
bullets([
    'Each item\'s stock quantity is increased by the quantity received.',
    'The item\'s buying price is updated to the latest price from this delivery.',
    'A stock movement record is created (type: "restock") referencing the GRN.',
    'If linked to an LPO, the purchase order status changes to "received".',
])
tip('Always create a GRN when receiving goods — this is the only way the system updates your stock from a purchase.')

doc.add_page_break()

# ============================================================
# 7. STOCK TAKE
# ============================================================
doc.add_heading('7. Stock Take (Physical Audit)', level=1)
doc.add_paragraph(
    'A stock take lets you compare what the system thinks you have against what is physically '
    'on your shelves. This catches theft, damage, miscounts, and system errors.'
)
doc.add_paragraph('Available to: Users with the "can_edit_stock" permission (Admins always have this).')

doc.add_heading('7.1 Starting a Stock Take', level=2)
steps([
    'Navigate to Stock Takes.',
    'Click "New Stock Take" or "Start Stock Take".',
    '(Optional) Add notes (e.g., "Monthly stock count — May 2026").',
    'Click "Create".',
])
doc.add_paragraph(
    'The system automatically creates a snapshot of every active item in your inventory, '
    'recording the expected quantity (what the system says you should have). '
    'The stock take is created in "Draft" status.'
)

doc.add_heading('7.2 Entering Physical Counts', level=2)
steps([
    'Open the stock take you just created.',
    'You will see a list of all items with their Expected Quantity.',
    'Walk through your pharmacy and count each item on the shelf.',
    'Enter the Actual Quantity for each item.',
    'The system automatically calculates the Variance (Actual − Expected).',
    'Click "Save" periodically to preserve your progress.',
])

add_table(
    ['Expected', 'Actual', 'Variance', 'Meaning'],
    [
        ['50', '50', '0', 'Stock matches perfectly ✓'],
        ['50', '45', '−5', '5 units unaccounted for (loss/theft/damage)'],
        ['50', '53', '+3', '3 extra units (possible receiving error)'],
    ]
)

doc.add_heading('7.3 Completing a Stock Take', level=2)
steps([
    'Once all items have been counted, review the variances.',
    'Click "Complete" or "Finalize Stock Take".',
    'Confirm when prompted.',
])
doc.add_heading('What happens:', level=3)
bullets([
    'The stock take status changes to "Completed" (cannot be edited further).',
    'For every item with a non-zero variance, the system forces the stock quantity to match the actual count.',
    'A stock movement record (type: "adjustment") is created for each corrected item.',
])
note('Completing a stock take permanently changes your stock levels. Make sure all counts are accurate before finalizing.')

doc.add_page_break()

# ============================================================
# 8. ANALYTICS
# ============================================================
doc.add_heading('8. Analytics & Reports', level=1)
doc.add_paragraph('Available to: Admin and Super Admin only')
doc.add_paragraph('The Analytics section provides business intelligence to help you make informed decisions.')

doc.add_heading('Dashboard Summary', level=2)
doc.add_paragraph('Shows at a glance:')
bullets([
    "Today's total sales, revenue, and profit.",
    'Period totals (filterable by date range — defaults to current month).',
    'Total active inventory items and categories.',
    'Number of low-stock items requiring attention.',
    'Number of items expiring soon.',
])

doc.add_heading('Sales Trend', level=2)
doc.add_paragraph('A chart showing your sales performance over time. Group by day, week, or month. Shows sales count, revenue, and profit per period.')

doc.add_heading('Top Selling Items', level=2)
doc.add_paragraph('Ranks your products by total revenue. Use this to identify your best performers and focus restocking on high-demand items.')

doc.add_heading('Category Breakdown', level=2)
doc.add_paragraph('Shows revenue and profit broken down by product category. Helps you understand which departments are most profitable.')

doc.add_heading('Item Report', level=2)
doc.add_paragraph('A detailed, line-by-line log of every item sold within a date range — including receipt number, date, quantity, buying price, selling price, and profit.')

doc.add_page_break()

# ============================================================
# 9. USER MANAGEMENT
# ============================================================
doc.add_heading('9. User Management', level=1)
doc.add_paragraph('Available to: Super Admin ONLY')
note('No other role can access these functions. The server will reject requests from Admin or Sales Attendant accounts.')

doc.add_heading('9.1 Creating a New User', level=2)
steps([
    'Navigate to User Management.',
    'Click "Add User" or "Register".',
    'Fill in: Username, Email, Full Name, Password.',
    'Select the Role: Sales Attendant, Admin, or Super Admin.',
    'Set Permissions (for Sales Attendants):',
    '   • Can Make Sales — allow/deny processing sales',
    '   • Can Receive Payments — allow/deny payment processing',
    '   • Can Edit Stock — allow/deny stock take participation',
    'Click "Create" or "Register".',
])

doc.add_heading('9.2 Editing Users', level=2)
doc.add_paragraph('Click on a user in the list to edit their details, role, or permissions.')

doc.add_heading('9.3 Disabling / Enabling Users', level=2)
doc.add_paragraph(
    'Click "Toggle Active" on a user to enable or disable their account. '
    'A disabled user cannot log in. Use this when staff leave or for temporary suspensions.'
)
tip('Disabling is preferred over deleting — it preserves the user\'s historical records (sales they processed, stock they received, etc.).')

doc.add_heading('9.4 Resetting Passwords', level=2)
doc.add_paragraph('Click "Reset Password" on a user to set a new password for them. The user does not need to provide their old password.')

doc.add_page_break()

# ============================================================
# 10. ROLE PERMISSIONS REFERENCE
# ============================================================
doc.add_heading('10. Role Permissions Reference', level=1)
doc.add_paragraph('Complete matrix of what each role can and cannot do:')

add_table(
    ['Function', 'Sales Attendant', 'Admin', 'Super Admin'],
    [
        ['Login / Logout', '✓', '✓', '✓'],
        ['Change own password', '✓', '✓', '✓'],
        ['Update own profile', '✓', '✓', '✓'],
        ['View inventory', '✓', '✓', '✓'],
        ['View stock alerts', '✓', '✓', '✓'],
        ['View stock movements', '✓', '✓', '✓'],
        ['Create a sale', '✓', '✓', '✓'],
        ['View sales history', '✓', '✓', '✓'],
        ['Print receipts', '✓', '✓', '✓'],
        ['View categories', '✓', '✓', '✓'],
        ['View suppliers / LPOs / GRNs', '✓', '✓', '✓'],
        ['View stock takes', '✓', '✓', '✓'],
        ['Add / Edit / Delete items', '✗', '✓', '✓'],
        ['Restock items', '✗', '✓', '✓'],
        ['Create / Edit categories', '✗', '✓', '✓'],
        ['Manage suppliers', '✗', '✓', '✓'],
        ['Create purchase orders', '✗', '✓', '✓'],
        ['Create GRNs (receive goods)', '✗', '✓', '✓'],
        ['Execute stock takes', '⚡', '✓', '✓'],
        ['View analytics & reports', '✗', '✓', '✓'],
        ['Register new users', '✗', '✗', '✓'],
        ['Edit / Delete users', '✗', '✗', '✓'],
        ['Enable / Disable accounts', '✗', '✗', '✓'],
        ['Reset user passwords', '✗', '✗', '✓'],
    ]
)
doc.add_paragraph('⚡ = Requires "can_edit_stock" permission to be explicitly granted by Super Admin')

doc.add_page_break()

# ============================================================
# 11. TROUBLESHOOTING
# ============================================================
doc.add_heading('11. Troubleshooting', level=1)

add_table(
    ['Problem', 'Possible Cause', 'Solution'],
    [
        ['Cannot log in', 'Wrong credentials or disabled account', 'Verify username/password. Contact Super Admin if account may be disabled.'],
        ['"Insufficient stock" error', 'Requested quantity exceeds available stock', 'Check current stock level. Restock or reduce sale quantity.'],
        ['"Insufficient permissions"', 'Your role lacks access to this function', 'Contact Admin or Super Admin to adjust your permissions.'],
        ['"Token expired"', 'Session has timed out', 'Log in again. This is normal after inactivity.'],
        ['Missing items in inventory', 'Item may have been soft-deleted', 'Contact Admin to check inactive items.'],
        ['Stock numbers seem wrong', 'Unrecorded sales or receiving errors', 'Run a Stock Take to reconcile physical vs system counts.'],
        ['"Too many requests" error', 'Rate limiting triggered', 'Wait a moment and try again. If persistent, contact support.'],
    ]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of User Guide —')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Moreran Chemist © 2026. All rights reserved.')

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Moreran_Chemist_User_Guide.docx')
doc.save(output_path)
print(f'User guide saved to: {output_path}')
